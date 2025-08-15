# Importaciones relativas del proyecto
from ..models import Universidad, Tema, Curso, Pregunta, UserProfile
from ..forms import Pregunta, FiltroPreguntaForm, PreguntaForm

# Django - shortcuts y decoradores
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.admin.views.decorators import staff_member_required
from django.views.decorators.http import require_POST
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.core.paginator import Paginator
from django.http import HttpResponse, FileResponse, Http404, JsonResponse
from django.conf import settings
from django.utils import timezone
from django.utils.text import slugify
from django.views.decorators.clickjacking import xframe_options_exempt
from django.contrib import messages
from django.urls import reverse

# Python estándar
import os
import io
import logging
from collections import defaultdict
from datetime import timedelta

# Librerías de terceros para manejo de documentos DOCX
from docx import Document
from docxcompose.composer import Composer

try:
    from docxcompose.composer import ImportFormatMode
except ImportError:
    ImportFormatMode = None

from docx.shared import Pt, Inches  # Tamaño de fuente y márgenes
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn

# Importación de vistas de autenticación propias
from .auth_views import exclude_supervisor


# Gestión de Preguntas
from django.core.paginator import Paginator

@exclude_supervisor
@login_required
def pregunta_list(request):
    user_profile = get_object_or_404(UserProfile, user=request.user)

    if request.user.is_superuser:
        qs = Pregunta.objects.filter(usuario=user_profile)
    else:
        limite = timezone.now() - timedelta(days=1)
        qs = Pregunta.objects.filter(usuario=user_profile, fecha_creacion__gte=limite)

    qs = qs.order_by('-fecha_creacion')

    universidad_id = request.GET.get('universidad')
    curso_id = request.GET.get('curso')
    tema_id = request.GET.get('tema')
    nivel = request.GET.get('nivel')
    tiempo_filtro = request.GET.get('tiempo_filtro')

    if tiempo_filtro:
        try:
            minutos = int(tiempo_filtro)
            limite_tiempo = timezone.now() - timedelta(minutes=minutos)
            qs = qs.filter(fecha_creacion__gte=limite_tiempo)
        except ValueError:
            pass

    if universidad_id:
        qs = qs.filter(universidad_id=universidad_id)
    if curso_id:
        qs = qs.filter(curso_id=curso_id)
    if tema_id:
        qs = qs.filter(tema_id=tema_id)
    if nivel:
        qs = qs.filter(nivel=nivel)

    paginator = Paginator(qs, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    form = FiltroPreguntaForm(request.GET or None)

    context = {
        'total_preguntas': Pregunta.objects.filter(usuario=user_profile).count(),
        'page_obj': page_obj,
        'form': form,
        'universidades': Universidad.objects.all(),
        'cursos_para_uni': Curso.objects.filter(universidades__id=universidad_id) if universidad_id else [],
        'temas_para_curso': Tema.objects.filter(curso_id=curso_id) if curso_id else [],
        'universidad_filter': universidad_id,
        'curso_filter': curso_id,
        'tema_filter': tema_id,
        'nivel_filter': nivel,
        'tiempo_filtro': tiempo_filtro,
        'now': timezone.now(),
    }

    return render(request, 'Preguntas/pregunta_list.html', context)


@login_required
def pregunta_list_supervisor(request):
    qs = Pregunta.objects.select_related('usuario__user', 'universidad', 'curso', 'tema').all()

    buscar_nombre = request.GET.get('nombre')
    if buscar_nombre:
        qs = qs.filter(nombre__icontains=buscar_nombre)

    buscar_usuario = request.GET.get('usuario')
    if buscar_usuario:
        qs = qs.filter(usuario_id=buscar_usuario)

    qs = qs.order_by('-fecha_creacion')

    paginator = Paginator(qs, 20)
    page = request.GET.get('page')
    qs_paginated = paginator.get_page(page)

    from django.contrib.auth.models import User
    usuarios = User.objects.all()

    context = {
        'preguntas': qs_paginated,
        'buscar_nombre': buscar_nombre,
        'usuarios': usuarios,
        'buscar_usuario': buscar_usuario,
    }

    return render(request, 'Preguntas/lista_supervisor.html', context)

@login_required
@exclude_supervisor
def pregunta_create(request):
    if request.method == 'POST':
        form = PreguntaForm(request.POST, request.FILES)
        if form.is_valid():
            pregunta = form.save(commit=False)
            user_profile = UserProfile.objects.get(user=request.user)
            pregunta.usuario = user_profile

            count = Pregunta.objects.filter(
                universidad=pregunta.universidad,
                curso=pregunta.curso,
                tema=pregunta.tema,
                nivel=pregunta.nivel
            ).count() + 1

            pregunta.nombre = f"{pregunta.universidad.id}{pregunta.curso.id}{pregunta.tema.id}{pregunta.nivel}{count}"
            pregunta.save()

            if 'contenido' in request.FILES:
                pregunta.contenido = request.FILES['contenido']
                pregunta.save()

            messages.success(request, 'Pregunta creada exitosamente.')

            # ✅ SIMULAR nuevo formulario
            data = request.POST.copy()
            data['nombre'] = ''
            nuevo_formulario = PreguntaForm(data, is_update=False)
            nuevo_formulario.fields['contenido'].required = False

            return render(request, 'Preguntas/pregunta_form.html', {
                'form': nuevo_formulario,
                'title': 'Nueva Pregunta'
            })

        else:
            return render(request, 'Preguntas/pregunta_form.html', {
                'form': form,
                'title': 'Nueva Pregunta'
            })
    else:
        form = PreguntaForm()

    return render(request, 'Preguntas/pregunta_form.html', {
        'form': form,
        'title': 'Nueva Pregunta'
    })

@login_required
@exclude_supervisor
def pregunta_update(request, pk):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        messages.error(request, 'No se encontró el perfil de usuario.')
        return redirect('pregunta-list')

    if request.user.is_superuser:
        pregunta = get_object_or_404(Pregunta, pk=pk)
    else:
        pregunta = get_object_or_404(Pregunta, pk=pk, usuario=user_profile)
        
    if request.method == 'POST':
        form = PreguntaForm(request.POST, request.FILES, instance=pregunta, is_update=True)
        
        if form.is_valid():
            nuevo_archivo = request.FILES.get('contenido')

            # Eliminar PDF existente si hay cambios
            pdf_dir = os.path.join(settings.MEDIA_ROOT, 'pdfs')
            safe_filename = sanitize_filename(pregunta.nombre)
            pdf_filename = f"{safe_filename}_{pregunta.id}.pdf"
            pdf_path = os.path.join(pdf_dir, pdf_filename)
            
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                    logger.info(f"PDF eliminado por actualización de pregunta: {pdf_path}")
                except Exception as e:
                    logger.error(f"No se pudo eliminar PDF: {e}")

            # Actualizar campos editables
            pregunta.respuesta = form.cleaned_data['respuesta']
            pregunta.nivel = form.cleaned_data['nivel']
            pregunta.tiene_solucion = form.cleaned_data['tiene_solucion']
            
            if nuevo_archivo:
                pregunta.contenido = nuevo_archivo
                logger.info(f"📝 Archivo actualizado para pregunta {pk}: {nuevo_archivo.name}")
            
            pregunta.save()

            messages.success(request, 'Pregunta actualizada con éxito.')
            return redirect('pregunta-list')
        else:
            # Mostrar errores del formulario
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"Error en {field}: {error}")
    else:
        form = PreguntaForm(instance=pregunta, is_update=True)

    return render(request, 'Preguntas/pregunta_form.html', {
        'form': form,
        'pregunta': pregunta,
        'title': 'Editar Pregunta',
        'is_update': True,
        'current_file': pregunta.contenido.name if pregunta.contenido else None
    })

@require_POST
@login_required
def actualizar_rapido_pregunta(request):
    try:
        pregunta_id = request.POST.get('id')
        tiene_solucion = request.POST.get('tiene_solucion', 'false') == 'true'
        alternativa = request.POST.get('alternativa', '').upper()

        logger.info(f"Actualización rápida solicitada - User: {request.user}, Pregunta ID: {pregunta_id}")

        if not pregunta_id:
            logger.warning("ID de pregunta no proporcionado")
            return JsonResponse({'success': False, 'error': 'ID de pregunta no proporcionado'}, status=400)

        if alternativa not in ['A', 'B', 'C', 'D', 'E']:
            logger.warning(f"Alternativa inválida recibida: {alternativa}")
            return JsonResponse({'success': False, 'error': 'Alternativa inválida'}, status=400)

        try:
            pregunta = Pregunta.objects.select_related('usuario__user').get(id=pregunta_id)
            
            # Verificar permisos mejorado
            if not (request.user.is_superuser or 
                   (hasattr(pregunta, 'usuario') and 
                    pregunta.usuario and 
                    pregunta.usuario.user == request.user)):
                logger.warning(f"Intento de edición no autorizado. User: {request.user}, Pregunta: {pregunta_id}")
                return JsonResponse(
                    {'success': False, 'error': 'No tienes permisos para editar esta pregunta'},
                    status=403
                )

            # Actualizar campos
            pregunta.tiene_solucion = tiene_solucion
            pregunta.respuesta = alternativa
            pregunta.save(update_fields=['tiene_solucion', 'respuesta'])

            logger.info(f"Pregunta {pregunta_id} actualizada correctamente por {request.user}")
            return JsonResponse({'success': True})
            
        except Pregunta.DoesNotExist:
            logger.error(f"Pregunta no encontrada: {pregunta_id}")
            return JsonResponse({'success': False, 'error': 'Pregunta no encontrada'}, status=404)
            
    except Exception as e:
        logger.error(f"Error en actualización rápida: {str(e)}", exc_info=True)
        return JsonResponse({'success': False, 'error': 'Error interno del servidor'}, status=500)
    
@login_required
@exclude_supervisor
def pregunta_delete(request, pk):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        messages.error(request, 'No se encontró el perfil de usuario.')
        return redirect('pregunta-list')

    if request.user.is_superuser:
        pregunta = get_object_or_404(Pregunta, pk=pk)
    else:
        pregunta = get_object_or_404(Pregunta, pk=pk, usuario=user_profile)

    # Capturamos la URL anterior solo para el GET (para mostrar "volver")
    referer = request.META.get('HTTP_REFERER', reverse('pregunta-list'))

    if request.method == 'POST':
        pregunta.delete()
        messages.success(request, 'Pregunta eliminada exitosamente.')
        return redirect('pregunta-list')  # Redirige siempre a la lista

    return render(request, 'Preguntas/pregunta_confirm_delete.html', {
        'pregunta': pregunta,
        'volver_url': referer, 
    })

@login_required
@exclude_supervisor
def eliminar_preguntas(request):
    if request.method == 'POST':
        pregunta_ids = request.POST.getlist('preguntas')
        
        try:
            user_profile = UserProfile.objects.get(user=request.user)
        except UserProfile.DoesNotExist:
            messages.error(request, "No se encontró el perfil de usuario.")
            return redirect('pregunta-list')

        # Verificar que todas las preguntas pertenecen al usuario (a menos que sea superuser)
        if request.user.is_superuser:
            preguntas = Pregunta.objects.filter(id__in=pregunta_ids)
        else:
            preguntas = Pregunta.objects.filter(id__in=pregunta_ids, usuario=user_profile)

        count = preguntas.count()
        if count == 0:
            messages.error(request, 'No se encontraron preguntas para eliminar.')
            return redirect('pregunta-list')

        # Eliminar las preguntas
        preguntas.delete()
        messages.success(request, f'Se eliminaron {count} pregunta(s) correctamente.')
        
        return redirect('pregunta-list')

    # Si no es POST, redirigir
    return redirect('pregunta-list')


#desde aquí empecé a modificar lo del formato de las preguntas
#para darle 2 columnas al doc final
def set_tres_columns(section):
    sectPr = section._sectPr  # Obtener el elemento de la sección
    cols = OxmlElement('w:cols')
    cols.set(ns.qn('w:num'), '3')  # Establecer dos columnas
    sectPr.append(cols)

def set_margenes(section):
    """Configura los márgenes del documento según lo solicitado."""
    section.top_margin = Inches(2 / 2.54)  # 2 cm
    section.left_margin = Inches(0.76 / 2.54)  # 0.76 cm
    section.right_margin = Inches(0.76 / 2.54)  # 0.76 cm
    section.bottom_margin = Inches(3.25 / 2.54)  # 3.25 cm

def aplicar_formato_texto(doc):
    """Aplica Arial Narrow y tamaño 9 pt a todo el contenido del documento."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial Narrow"
            run.font.size = Pt(9)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "Arial Narrow")
    for style in doc.styles:
        if style.type == 1:  # Solo afecta estilos de párrafo
            if style.name.lower() in ["list paragraph", "lista numerada", "lista con viñetas"]:
                style.font.name = "Arial Narrow"
                style.font.size = Pt(9)

def eliminar_ultimo_parrafo_si_vacio(document):
    if len(document.paragraphs) == 0:
        return

    last_paragraph = document.paragraphs[-1]
    if not last_paragraph.text.strip():
        p_element = last_paragraph._element
        p_element.getparent().remove(p_element)

def combinar_documentos(preguntas):
    """Combina documentos mostrando solo nombres de preguntas sin espacios adicionales"""
    master_doc = Document()
    composer = Composer(master_doc)
    
    # Configurar documento con 3 columnas y márgenes
    set_tres_columns(master_doc.sections[0])
    set_margenes(master_doc.sections[0])
    
    # Eliminar espaciado por defecto en el documento base
    style = master_doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(9)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0  # Espaciado simple

    # Agrupar preguntas por curso y tema (solo para orden)
    preguntas_ordenadas = defaultdict(lambda: defaultdict(list))
    for pregunta in preguntas:
        if pregunta.contenido and hasattr(pregunta.contenido, 'path'):
            preguntas_ordenadas[pregunta.curso.nombre][pregunta.tema.nombre].append(pregunta)

    # Procesar preguntas
    for curso, temas in sorted(preguntas_ordenadas.items()):
        for tema, preguntas_tema in sorted(temas.items()):
            for pregunta in preguntas_tema:
                try:
                    # Agregar nombre de pregunta sin espaciado
                    p = master_doc.add_paragraph(style='Normal')
                    run = p.add_run(f"Pregunta: {pregunta.nombre}")
                    run.bold = True
                    
                    # Procesar contenido de la pregunta
                    sub_doc = Document(pregunta.contenido.path)
                    
                    # Eliminar configuraciones de sección y ajustar formato
                    for element in sub_doc.element.body:
                        if element.tag.endswith('sectPr'):
                            sub_doc.element.body.remove(element)
                    
                    # Aplicar formato sin espacios a todos los párrafos
                    for para in sub_doc.paragraphs:
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        for run in para.runs:
                            run.font.name = "Arial Narrow"
                            run.font.size = Pt(9)
                    
                    # Combinar documentos
                    if ImportFormatMode is not None:
                        composer.append(sub_doc, import_format=ImportFormatMode.KEEP_SOURCE_FORMATTING)
                    else:
                        composer.append(sub_doc)
                        
                except Exception as e:
                    logger.error(f"Error procesando pregunta {pregunta.id}: {str(e)}")
                    continue

    # Eliminar posibles párrafos vacíos finales
    eliminar_ultimo_parrafo_si_vacio(composer.doc)

    buffer = io.BytesIO()
    composer.save(buffer)
    buffer.seek(0)
    return buffer

logger = logging.getLogger(__name__)

def sanitize_filename(filename):
    """Sanitiza el nombre del archivo para evitar problemas"""
    # Remover caracteres problemáticos y limitar longitud
    filename = slugify(filename, allow_unicode=False)
    if len(filename) > 100:
        filename = filename[:100]
    return filename or "documento"

@xframe_options_exempt
@login_required
def vista_previa(request, pk):
    """Vista previa con reconversión solo si el PDF no existe"""
    try:
        pregunta = Pregunta.objects.get(pk=pk)
    except Pregunta.DoesNotExist:
        logger.warning(f"Pregunta {pk} no encontrada")
        raise Http404("La pregunta no existe")

    # Validar que el archivo DOCX existe
    docx_path = os.path.join(settings.MEDIA_ROOT, str(pregunta.contenido))
    if not os.path.exists(docx_path):
        logger.error(f"Archivo DOCX no encontrado: {docx_path}")
        raise Http404("El archivo DOCX no existe")

    # Crear directorio para PDFs si no existe
    pdf_dir = os.path.join(settings.MEDIA_ROOT, 'pdfs')
    os.makedirs(pdf_dir, exist_ok=True)

    # Generar nombre único para el PDF basado en la pregunta
    safe_filename = sanitize_filename(pregunta.nombre)
    pdf_filename = f"{safe_filename}_{pregunta.id}.pdf"
    pdf_path = os.path.join(pdf_dir, pdf_filename)

    # CONVERSIÓN SOLO SI EL PDF NO EXISTE
    if not os.path.exists(pdf_path):
        logger.info(f"Iniciando conversión: {docx_path} -> {pdf_path}")
        
        try:
            # Conversión con Aspose Words
            import aspose.words as aw
            
            # Cargar documento DOCX
            doc = aw.Document(docx_path)
            
            # Configurar opciones de guardado
            save_options = aw.saving.PdfSaveOptions()
            save_options.compliance = aw.saving.PdfCompliance.PDF17
            save_options.preserve_form_fields = True
            save_options.jpeg_quality = 90
            
            # Convertir a PDF
            doc.save(pdf_path, save_options)
            
            logger.info(f"✅ Conversión exitosa para pregunta {pk}")
            
        except ImportError:
            logger.error("❌ Aspose Words no está instalado")
            raise Http404("Error: Aspose Words no está disponible en el sistema")
        except Exception as e:
            logger.error(f"❌ Error en conversión: {e}")
            # Limpiar archivo parcial si existe
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                except:
                    pass
            raise Http404(f"Error al convertir archivo: {str(e)}")
    else:
        logger.info(f"✅ Usando PDF existente para pregunta {pk}")

    # VALIDACIONES FINALES
    if not os.path.exists(pdf_path):
        logger.error(f"❌ PDF no se generó correctamente: {pdf_path}")
        raise Http404("No se pudo generar el archivo PDF")

    try:
        file_size = os.path.getsize(pdf_path)
        if file_size == 0:
            logger.error(f"❌ PDF generado está vacío: {pdf_path}")
            os.remove(pdf_path)
            raise Http404("El archivo PDF generado está vacío")
        elif file_size < 100:
            logger.warning(f"⚠️ PDF sospechosamente pequeño ({file_size} bytes): {pdf_path}")
        
        logger.info(f"📄 Sirviendo PDF: {pdf_path} ({file_size} bytes)")
        
    except OSError as e:
        logger.error(f"❌ Error verificando PDF: {e}")
        raise Http404("Error al verificar el archivo PDF")

    return serve_pdf_file(pdf_path, pregunta.nombre)

def serve_pdf_file(pdf_path, original_name):
    try:
        pdf_file = open(pdf_path, 'rb')
        safe_download_name = sanitize_filename(original_name)
        
        # Obtener timestamp de modificación del archivo
        last_modified = int(os.path.getmtime(pdf_path))
        download_filename = f"{safe_download_name}_{last_modified}.pdf"
        
        response = FileResponse(
            pdf_file,
            content_type='application/pdf',
            filename=download_filename
        )
        
        # Headers para evitar cache
        response['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        response['X-Accel-Expires'] = '0'
        
        return response
    except Exception as e:
        logger.error(f"Error al servir PDF: {e}")
        raise Http404("Error al acceder al archivo PDF")

def cleanup_old_pdfs():
    """Limpia PDFs antiguos para ahorrar espacio"""
    try:
        pdf_dir = os.path.join(settings.MEDIA_ROOT, 'pdfs')
        if not os.path.exists(pdf_dir):
            return
        
        import time
        current_time = time.time()
        week_ago = current_time - (7 * 24 * 60 * 60)  # 7 días
        
        for filename in os.listdir(pdf_dir):
            file_path = os.path.join(pdf_dir, filename)
            if os.path.isfile(file_path):
                file_mtime = os.path.getmtime(file_path)
                if file_mtime < week_ago:
                    try:
                        os.remove(file_path)
                        logger.info(f"PDF antiguo eliminado: {filename}")
                    except Exception as e:
                        logger.warning(f"No se pudo eliminar {filename}: {e}")
                        
    except Exception as e:
        logger.error(f"Error en limpieza de PDFs: {e}")

@staff_member_required
@exclude_supervisor
def todas_las_preguntas(request):
    # QuerySet base
    qs = Pregunta.objects.all().order_by('-fecha_creacion')

    # Filtros
    buscar_nombre = request.GET.get('nombre')
    usuario_id = request.GET.get('usuario')
    fecha_filtro = request.GET.get('fecha')
    per_page = int(request.GET.get('per_page', 20))

    # Aplicar filtros
    if buscar_nombre:
        qs = qs.filter(nombre__icontains=buscar_nombre)
    
    if usuario_id:
        qs = qs.filter(usuario_id=usuario_id)
    
    if fecha_filtro:
        ahora = timezone.now()

        if fecha_filtro == 'hoy':
            limite = ahora - timedelta(days=1)
            qs = qs.filter(fecha_creacion__gte=limite)
        elif fecha_filtro == 'semana':
            limite = ahora - timedelta(days=7)
            qs = qs.filter(fecha_creacion__gte=limite)
        elif fecha_filtro == 'mes':
            limite = ahora - timedelta(days=30)
            qs = qs.filter(fecha_creacion__gte=limite)

    # Paginación
    paginator = Paginator(qs, per_page)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Contexto
    context = {
        'total_preguntas': qs.count(),
        'preguntas': page_obj,
        'usuarios': User.objects.all().order_by('username'),
        'now': timezone.now(),
        'modo_admin': True,
        'request': request,  # Para acceder a los parámetros GET en el template
    }
    
    # Respuesta completa
    return render(request, 'Preguntas/todas_las_preguntas.html', context)

@login_required
@exclude_supervisor
def descargar_preguntas(request):
    pregunta_ids = request.POST.getlist('preguntas')
    
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        messages.error(request, "No se encontró el perfil de usuario.")
        return redirect('pregunta-list')

    preguntas = Pregunta.objects.filter(id__in=pregunta_ids, usuario=user_profile)
    
    if not preguntas:
        messages.error(request, 'No se encontraron preguntas para descargar.')
        return redirect('pregunta-list')
    
    buffer = combinar_documentos(preguntas)
    
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="preguntas_combinadas.docx"'
    
    return response
