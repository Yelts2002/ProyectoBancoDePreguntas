import os
import tempfile
from contextlib import contextmanager
from io import BytesIO
from copy import deepcopy
from docx import Document
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib import messages
from django.shortcuts import render, redirect
from django.core.files import File
from ..forms import *
from .auth_views import exclude_supervisor
from ..models import Pregunta, UserProfile

@contextmanager
def temp_docx_file(content_bytes, suffix='.docx'):
    """Context manager mejorado para archivos temporales"""
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(content_bytes)
            temp_path = temp_file.name
        yield temp_path
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except:
                pass

def copy_related_parts(new_doc, original_doc, block_elements):
    """
    Copia las partes relacionadas (imágenes, gráficos) que son referenciadas
    en los elementos del bloque al nuevo documento.
    """
    # Obtener todas las relaciones del documento original
    original_rels = original_doc.part.rels
    
    # Identificar las relaciones utilizadas en el bloque
    used_rIds = set()
    for element in block_elements:
        # Buscar referencias a relaciones en el elemento
        for rel_attr in ['r:embed', 'r:link', 'v:imagedata']:
            for el in element.xpath(f'.//@{rel_attr}'):
                used_rIds.add(el)
    
    # Copiar las relaciones utilizadas al nuevo documento
    for rId in used_rIds:
        if rId in original_rels:
            rel = original_rels[rId]
            if rel.is_external:
                new_doc.part.rels.add_relationship(
                    rel.reltype,
                    rel._target,
                    rel.rId,
                    is_external=True
                )
            else:
                new_doc.part.rels.add_relationship(
                    rel.reltype,
                    rel.target_part,
                    rel.rId
                )

def create_exact_copy_docx(original_doc, block_elements):
    """
    Crea un nuevo documento que conserva exactamente todo el contenido
    incluyendo imágenes, tablas, ecuaciones, con sus formatos originales.
    Versión corregida del error 'element has no setter'.
    """
    # Crear un documento completamente nuevo
    new_doc = Document()
    
    # 1. Copiar los estilos del documento original (manera correcta)
    for style in original_doc.styles:
        try:
            # Solo copiamos los estilos que no existen en el nuevo documento
            if style.name not in new_doc.styles:
                new_style = new_doc.styles.add_style(style.name, style.type)
                # Copiar propiedades básicas del estilo
                if style.font:
                    new_style.font.name = style.font.name
                    new_style.font.size = style.font.size
                    new_style.font.bold = style.font.bold
                    new_style.font.italic = style.font.italic
                    new_style.font.underline = style.font.underline
                if hasattr(style, 'paragraph_format'):
                    new_style.paragraph_format.alignment = style.paragraph_format.alignment
                    new_style.paragraph_format.left_indent = style.paragraph_format.left_indent
                    new_style.paragraph_format.right_indent = style.paragraph_format.right_indent
        except Exception as e:
            print(f"Warning: No se pudo copiar el estilo {style.name}: {str(e)}")
            continue
    
    # 2. Configurar propiedades del documento
    new_doc._element.body.clear_content()
    
    # 3. Copiar las secciones (configuración de página)
    if original_doc.sections:
        new_sect = new_doc.sections[0]
        orig_sect = original_doc.sections[0]
        new_sect.start_type = orig_sect.start_type
        new_sect.orientation = orig_sect.orientation
        new_sect.page_width = orig_sect.page_width
        new_sect.page_height = orig_sect.page_height
        new_sect.left_margin = orig_sect.left_margin
        new_sect.right_margin = orig_sect.right_margin
        new_sect.top_margin = orig_sect.top_margin
        new_sect.bottom_margin = orig_sect.bottom_margin
        new_sect.header_distance = orig_sect.header_distance
        new_sect.footer_distance = orig_sect.footer_distance
    
    # Copiar elementos del bloque
    new_body = new_doc._element.body
    for element in block_elements:
        new_body.append(deepcopy(element))
    
    # Copiar partes relacionadas (imágenes, gráficos)
    copy_related_parts(new_doc, original_doc, block_elements)
    
    return new_doc

@login_required
@staff_member_required
@exclude_supervisor
def masivo_pregunta_create(request):
    if request.method == 'POST':
        form = CargaMasivaPreguntaForm(request.POST, request.FILES)
        if not form.is_valid():
            messages.error(request, 'Por favor corrija los errores en el formulario.')
            return render(request, 'Preguntas/masivo_pregunta_form.html', {
                'form': form,
                'title': 'Carga Masiva de Preguntas'
            })

        try:
            # Datos del formulario
            data = form.cleaned_data
            universidad = data['universidad']
            curso = data['curso']
            tema = data['tema']
            nivel = data['nivel']
            respuesta_default = data['respuesta_default']
            archivo_word = data['archivo']

            # Leer el archivo una sola vez
            archivo_bytes = archivo_word.read()
            user_profile = UserProfile.objects.get(user=request.user)
            
            # Procesar documento original para extraer bloques
            with temp_docx_file(archivo_bytes) as temp_path:
                doc = Document(temp_path)
                
                # Separar bloques por el separador '*****'
                all_blocks = []
                current_block = []
                
                for element in doc.element.body:
                    text_elements = element.xpath('.//w:t')
                    text = ''.join(t.text for t in text_elements if t.text).strip()

                    if text == '*****':
                        if current_block:
                            all_blocks.append(current_block)
                            current_block = []
                    else:
                        current_block.append(element)

                if current_block:
                    all_blocks.append(current_block)

            # Contador base para nombres de preguntas
            base_count = Pregunta.objects.filter(
                universidad=universidad,
                curso=curso,
                tema=tema,
                nivel=nivel
            ).count()

            preguntas_creadas = 0
            
            # Procesar cada bloque como pregunta
            for i, block_elements in enumerate(all_blocks, start=1):
                # Crear nuevo documento con todo el contenido original
                with temp_docx_file(archivo_bytes) as temp_path:
                    original_doc = Document(temp_path)
                    new_doc = create_exact_copy_docx(original_doc, block_elements)
                    
                    # Guardar en memoria
                    buffer = BytesIO()
                    new_doc.save(buffer)
                    buffer.seek(0)

                # Crear la pregunta
                pregunta = Pregunta(
                    universidad=universidad,
                    curso=curso,
                    tema=tema,
                    nivel=nivel,
                    respuesta=respuesta_default,
                    usuario=user_profile,
                    tiene_solucion=False,
                    nombre=f"{universidad.id}{curso.id}{tema.id}{nivel}{base_count + i}"
                )
                pregunta.save()
                
                # Guardar archivo con todo el contenido
                filename = f"pregunta_{pregunta.nombre}.docx"
                pregunta.contenido.save(filename, File(buffer))
                preguntas_creadas += 1

            messages.success(request, f'Se crearon {preguntas_creadas} preguntas exitosamente.')
            return redirect('pregunta_list')

        except Exception as e:
            messages.error(request, f'Error al procesar el archivo: {str(e)}')
            return render(request, 'Preguntas/masivo_pregunta_form.html', {
                'form': form,
                'title': 'Carga Masiva de Preguntas'
            })

    # GET request
    form = CargaMasivaPreguntaForm()
    return render(request, 'Preguntas/masivo_pregunta_form.html', {
        'form': form,
        'title': 'Carga Masiva de Preguntas'
    })